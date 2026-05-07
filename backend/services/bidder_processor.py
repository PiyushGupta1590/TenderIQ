"""
Stage 2 — Bidder Document Processor (Production-Grade)
=======================================================
PIPELINE:
1. File Detection & Parsing
2. Layout & Structure Analysis
3. Smart Chunking
4. Hybrid Extraction (Regex + LLM)
5. Table Extraction
6. Evidence Grounding
7. Cross-Verification
8. Confidence Scoring
9. Criteria Evaluation
10. Output Formatting
"""
from __future__ import annotations

import re
import uuid
import logging
from concurrent.futures import ThreadPoolExecutor, as_completed
from datetime import datetime, timezone
from pathlib import Path
from typing import List, Dict, Optional, Any
from dataclasses import asdict

from backend.models.bidder import BidderData, ExtractedField, ProcessingStatus
from backend.services.layout_analyzer import (
    analyze_pdf_layout,
    create_smart_chunks,
    extract_tables_from_pdf,
    DocumentBlock,
    TableBlock,
    BlockType
)
from backend.services.hybrid_extractor import (
    HybridExtractor,
    ExtractionResult,
    CriterionEvaluation,
    evaluate_criterion,
    Evidence,
    ExtractionStatus
)
from backend.services.ocr_engine import ocr_file, TESSERACT_AVAILABLE
from backend.config import DATA_DIR

# Optional docx support
try:
    from docx import Document as DocxDocument
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

logger = logging.getLogger(__name__)




_FIELD_PATTERNS = {
    "annual_turnover": [
        re.compile(
            r"(?:turnover|revenue|sales)[^:\n]*[:\s]+(?:Rs\.?|INR|₹|\$|USD)?\s*([\d,]+(?:\.\d+)?)\s*(crore|lakh|million|Cr|L|M)?",
            re.IGNORECASE
        ),
        re.compile(
            r"(?:total\s+)?(?:annual\s+)?(?:revenue|income)[^:\n]*[:\s]+(?:Rs\.?|INR|₹|\$)?\s*([\d,]+(?:\.\d+)?)\s*(crore|lakh|Cr|L)?",
            re.IGNORECASE
        ),
    ],
    "net_worth": [
        re.compile(
            r"net\s*worth[^:\n]*[:\s]+(?:Rs\.?|INR|₹|\$)?\s*([\d,]+(?:\.\d+)?)\s*(crore|lakh|Cr|L)?",
            re.IGNORECASE
        ),
    ],
    "bid_security": [
        re.compile(
            r"(?:bid\s+security|EMD|earnest\s+money)[^:\n]*[:\s]+(?:Rs\.?|INR|₹|\$)?\s*([\d,]+(?:\.\d+)?)\s*(crore|lakh|Cr|L)?",
            re.IGNORECASE
        ),
    ],
    "performance_guarantee": [
        re.compile(
            r"performance\s+(?:guarantee|security|bond)[^:\n]*[:\s]+(?:Rs\.?|INR|₹|\$)?\s*([\d,]+(?:\.\d+)?)\s*(crore|lakh|Cr|L|%)?",
            re.IGNORECASE
        ),
    ],
    "experience_years": [
        re.compile(r"(\d+)\s+years?\s+(?:of\s+)?experience", re.IGNORECASE),
        re.compile(r"experience[^:\n]*[:\s]+(\d+)\s+years?", re.IGNORECASE),
    ],
    "completed_projects": [
        re.compile(r"(?:completed|executed|delivered)\s+(\d+)\s+(?:similar|projects?|works?|contracts?)", re.IGNORECASE),
        re.compile(r"(\d+)\s+(?:projects?|works?|contracts?)\s+(?:completed|executed)", re.IGNORECASE),
    ],
    "similar_work_value": [
        re.compile(
            r"(?:single\s+)?(?:similar\s+work|work\s+order)[^:\n]*[:\s]+(?:Rs\.?|INR|₹)?\s*([\d,]+(?:\.\d+)?)\s*(crore|lakh|Cr|L)?",
            re.IGNORECASE
        ),
    ],
    "security_force_experience": [
        re.compile(
            r"(?:experience|worked?|served?|supplied?)\s+(?:with|for|to)\s+(?:CRPF|CISF|BSF|ITBP|police|paramilitary|defence|armed\s+forces)",
            re.IGNORECASE
        ),
        re.compile(r"(?:CRPF|CISF|BSF|ITBP|NIA|NSG)\s+(?:contract|order|work|project)", re.IGNORECASE),
    ],
    "warranty_period": [
        re.compile(r"(?:warranty|guarantee)\s+(?:of|period)[:\s]+(\d+)\s+(?:year|month)", re.IGNORECASE),
        re.compile(r"(\d+)\s+(?:year|month)\s+(?:warranty|guarantee)", re.IGNORECASE),
    ],
    "iso_certification": [
        re.compile(r"ISO\s*(9001|14001|27001|45001|22000)[:\s]*(?::\s*\d{4})?", re.IGNORECASE),
        re.compile(r"certified\s+(?:to|under|for)\s+ISO", re.IGNORECASE),
    ],
    "gst_registration": [
        re.compile(
            r"GST(?:IN)?\s*(?:No\.?|Number|Registration)?[:\s]*([0-9]{2}[A-Z]{5}[0-9]{4}[A-Z][0-9A-Z]Z[0-9A-Z])",
            re.IGNORECASE
        ),
        re.compile(r"GST\s+(?:Registration|Enrolled|Registered|Certificate)", re.IGNORECASE),
        re.compile(r"Goods\s+and\s+Services\s+Tax\s+(?:number|registration|certificate)", re.IGNORECASE),
    ],
    "pan_card": [
        re.compile(r"PAN[:\s]*([A-Z]{5}[0-9]{4}[A-Z])", re.IGNORECASE),
        re.compile(r"Permanent\s+Account\s+Number[:\s]*([A-Z]{5}[0-9]{4}[A-Z])?", re.IGNORECASE),
        re.compile(r"PAN\s+(?:Card|Certificate|Number)", re.IGNORECASE),
    ],
    "msme_registration": [
        re.compile(r"(?:MSME|Udyam|Udyog\s+Aadhar|SSI)\s+(?:Registration|Certificate|No\.?)[:\s]*([A-Z0-9\-]+)?", re.IGNORECASE),
        re.compile(r"Udyam\s+Registration\s+(?:Number|Certificate)", re.IGNORECASE),
    ],
    "labour_compliance": [
        re.compile(r"(?:EPF|PF)\s+(?:Registration|No\.?)[:\s]*([A-Z0-9/\-]+)?", re.IGNORECASE),
        re.compile(r"ESIC\s+(?:Registration|No\.?|Code)[:\s]*([0-9]+)?", re.IGNORECASE),
        re.compile(r"(?:provident\s+fund|labour\s+license)\s+(?:registration|number)", re.IGNORECASE),
    ],
    "registration": [
        re.compile(r"registered\s+(?:with|under|in)\s+([^,.\n]+)", re.IGNORECASE),
        re.compile(r"registration\s+(?:no\.?|number)[:\s]*([A-Z0-9\-/]+)", re.IGNORECASE),
    ],
    "company_registration": [
        re.compile(r"(?:CIN|Company\s+Identification\s+Number)[:\s]*([A-Z0-9]+)?", re.IGNORECASE),
        re.compile(r"(?:certificate\s+of\s+incorporation|ROC\s+certificate)", re.IGNORECASE),
    ],
    "blacklist_free": [
        re.compile(r"(?:not|never)\s+(?:blacklisted|debarred|banned|suspended)", re.IGNORECASE),
        re.compile(r"(?:blacklist|debarment)\s+(?:free|clear|clean)", re.IGNORECASE),
    ],
    "security_clearance": [
        re.compile(r"(?:security\s+clearance|police\s+verification)\s+(?:certificate|report|no\.?)[:\s]*([A-Z0-9/\-]+)?", re.IGNORECASE),
        re.compile(r"(?:security\s+clearance|antecedent\s+verification)\s+(?:obtained|issued|granted)", re.IGNORECASE),
    ],
    "solvency_certificate": [
        re.compile(r"(?:solvency|financial\s+soundness)\s+certificate", re.IGNORECASE),
        re.compile(r"credit\s+rating\s+(?:certificate|report|grade)", re.IGNORECASE),
    ],
    "key_personnel": [
        re.compile(r"(\d+)\s+(?:qualified|experienced|technical)?\s*(?:personnel|staff|engineers?)", re.IGNORECASE),
    ],
}

_UNIT_MAP = {
    "crore": 1e7, "cr": 1e7,
    "lakh": 1e5, "l": 1e5,
    "million": 1e6, "m": 1e6,
}

# Fields where we just confirm presence (boolean)
_PRESENCE_FIELDS = {
    "iso_certification", "gst_registration", "pan_card", "msme_registration",
    "labour_compliance", "registration", "company_registration",
    "blacklist_free", "security_clearance", "solvency_certificate",
    "security_force_experience", "power_of_attorney", "litigation_free",
}

# Fields where a raw match string is sufficient (e.g. cert number)
_STRING_EXTRACT_FIELDS = {
    "gst_registration", "pan_card", "msme_registration", "labour_compliance",
    "registration", "company_registration", "security_clearance",
}


# ── Helper functions for DOCX ────────────────────────────────────────────────

def _read_docx_text(docx_path: str) -> Tuple[List[str], List[float]]:
    """Read text from DOCX file, return list of page texts and confidences"""
    if not DOCX_AVAILABLE:
        return [], []
    try:
        doc = DocxDocument(docx_path)
        paragraphs = []
        for para in doc.paragraphs:
            if para.text.strip():
                paragraphs.append(para.text)
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip():
                        paragraphs.append(cell.text)
        full_text = "\n".join(paragraphs)
        return [full_text], [0.95]   # DOCX text is native, high confidence
    except Exception as e:
        logger.error(f"Error reading DOCX: {e}")
        return [], []


# ── Main Processor Class ─────────────────────────────────────────────────────

class BidderProcessor:
    """Production-grade bidder document processor"""
    
    def __init__(self, tender_criteria: Optional[List[Dict]] = None):
        """
        Initialize processor with tender criteria
        
        Args:
            tender_criteria: List of tender rules/criteria to evaluate against
        """
        self.tender_criteria = tender_criteria or []
        self.hybrid_extractor = HybridExtractor(_FIELD_PATTERNS)
    
    def process_bidder_files(
        self,
        file_paths: List[str],
        bidder_name: str,
        tender_id: str
    ) -> Dict[str, Any]:
        """
        Process bidder files with production-grade pipeline
        
        Args:
            file_paths: List of file paths to process
            bidder_name: Bidder company name
            tender_id: Associated tender ID
        
        Returns:
            Dictionary with bidder data and criteria evaluations
        """
        bidder_id = str(uuid.uuid4())
        now = datetime.now(timezone.utc).isoformat()
        
        processing_notes = []
        all_blocks = []
        all_chunks = []
        all_tables = []
        
        # STAGE 1: File Detection & Parsing
        logger.info(f"Processing {len(file_paths)} files for bidder: {bidder_name}")
        
        for file_path in file_paths:
            ext = Path(file_path).suffix.lower()
            fname = Path(file_path).name
            
            try:
                if ext == ".pdf":
                    # Run layout analysis and table extraction concurrently
                    # (both open the same PDF, but I/O overlap saves time)
                    with ThreadPoolExecutor(max_workers=2) as pdf_ex:
                        f_blocks = pdf_ex.submit(analyze_pdf_layout, file_path)
                        f_tables = pdf_ex.submit(extract_tables_from_pdf, file_path)
                        blocks = f_blocks.result()
                        tables = f_tables.result()

                    all_blocks.extend(blocks)
                    all_tables.extend(tables)
                    logger.info(f"Extracted {len(blocks)} blocks and {len(tables)} tables from {fname}")
                
                elif ext in (".docx", ".doc"):
                    # DOCX processing
                    page_texts, _ = _read_docx_text(file_path)
                    # Convert to blocks
                    for page_num, text in enumerate(page_texts, start=1):
                        all_blocks.append(DocumentBlock(
                            block_type=BlockType.PARAGRAPH,
                            content=text,
                            page_number=page_num,
                            section_name=None,
                            confidence=0.95
                        ))
                
                elif ext in (".png", ".jpg", ".jpeg", ".tiff", ".bmp", ".webp"):
                    if TESSERACT_AVAILABLE:
                        text, conf = ocr_file(file_path)
                        all_blocks.append(DocumentBlock(
                            block_type=BlockType.PARAGRAPH,
                            content=text,
                            page_number=1,
                            section_name=None,
                            confidence=conf
                        ))
                        
                        if conf < 0.60:
                            processing_notes.append(
                                f"Low OCR confidence ({conf:.0%}) for {fname} - manual review recommended"
                            )
                    else:
                        processing_notes.append(f"Tesseract not available - skipped {fname}")
                
                else:
                    processing_notes.append(f"Unsupported format: {fname}")
            
            except Exception as e:
                logger.error(f"Error processing {fname}: {e}")
                processing_notes.append(f"Error processing {fname}: {str(e)}")
        
        # STAGE 3: Smart Chunking
        if all_blocks:
            all_chunks = create_smart_chunks(all_blocks, max_chunk_size=2000)
            logger.info(f"Created {len(all_chunks)} smart chunks")
        
        # STAGE 4 & 5: Hybrid Extraction
        extraction_results = self._extract_hybrid(all_blocks, all_chunks, all_tables)
        
        # STAGE 6: Criteria Evaluation (if criteria provided)
        criteria_evaluations = []
        if self.tender_criteria:
            criteria_evaluations = self._evaluate_criteria(extraction_results)
        
        # STAGE 7: Calculate aggregate confidence
        confidences = [r.confidence for r in extraction_results if r.confidence > 0]
        avg_confidence = sum(confidences) / len(confidences) if confidences else 0.0
        
        # STAGE 8: Determine status
        has_conflicts = any(r.status == ExtractionStatus.CONFLICT for r in extraction_results)
        has_missing = any(r.status == ExtractionStatus.MISSING for r in extraction_results)
        has_low_conf = any(r.status == ExtractionStatus.LOW_CONFIDENCE for r in extraction_results)
        
        if has_conflicts:
            status = ProcessingStatus.MANUAL_REVIEW
            processing_notes.append("Conflicts detected between extraction methods - review required")
        elif has_missing and len([r for r in extraction_results if r.value is not None]) < 5:
            status = ProcessingStatus.MANUAL_REVIEW
            processing_notes.append("Multiple required fields missing")
        elif has_low_conf:
            status = ProcessingStatus.MANUAL_REVIEW
            processing_notes.append("Low confidence extractions detected")
        else:
            status = ProcessingStatus.COMPLETE
        
        # STAGE 9: Format output
        output = {
            "bidder_id": bidder_id,
            "bidder_name": bidder_name,
            "tender_id": tender_id,
            "uploaded_files": [Path(p).name for p in file_paths],
            "uploaded_at": now,
            "status": status.value,
            "progress": 100,
            "processing_notes": processing_notes,
            
            # Structured extraction results
            "extraction_results": [self._format_extraction_result(r) for r in extraction_results],
            
            # Criteria evaluations
            "criteria_results": [self._format_criterion_evaluation(c) for c in criteria_evaluations],
            
            # Metrics
            "avg_confidence": round(avg_confidence, 2),
            "total_fields_extracted": len(extraction_results),
            "fields_with_conflicts": sum(1 for r in extraction_results if r.status == ExtractionStatus.CONFLICT),
            "fields_missing": sum(1 for r in extraction_results if r.status == ExtractionStatus.MISSING),
            
            # Legacy format for backward compatibility
            "extracted_fields": self._convert_to_legacy_format(extraction_results),
            "ocr_confidence_avg": avg_confidence,
            "confirmed": False,
        }
        
        return output
    
    def _extract_hybrid(
        self,
        blocks: List[DocumentBlock],
        chunks: List[Dict],
        tables: List[TableBlock]
    ) -> List[ExtractionResult]:
        """Hybrid extraction combining regex, LLM, and table extraction.
        All field extractions run in parallel threads.
        """
        fields = list(_FIELD_PATTERNS.keys())

        def _extract_one(field: str) -> ExtractionResult:
            return self.hybrid_extractor.extract_field_hybrid(
                field=field,
                blocks=blocks,
                chunks=chunks,
                context_hint=f"Extract {field.replace('_', ' ')} from bidder document"
            )

        # Run all field extractions concurrently (LLM calls overlap)
        max_workers = min(8, len(fields))
        with ThreadPoolExecutor(max_workers=max_workers) as executor:
            results: List[ExtractionResult] = list(executor.map(_extract_one, fields))

        # Extract from tables (if relevant fields found)
        table_fields = self._extract_from_tables(tables)
        for field, value, evidence in table_fields:
            existing = next((r for r in results if r.field == field), None)
            if existing and existing.value is None:
                existing.value = value
                existing.evidence = evidence
                existing.extraction_method = "table"
                existing.status = ExtractionStatus.OK

        return results
    
    def _extract_from_tables(self, tables: List[TableBlock]) -> List[tuple]:
        """Extract structured data from tables"""
        extractions = []
        
        for table in tables:
            # Look for financial data in tables
            for row_idx, row in enumerate(table.rows):
                row_text = " ".join(row).lower()
                
                # Check for known fields
                if "turnover" in row_text or "revenue" in row_text:
                    # Try to extract value from row
                    for cell in row:
                        if re.search(r'[\d,]+', cell):
                            try:
                                value = float(cell.replace(',', ''))
                                evidence = Evidence(
                                    text=f"Table row: {' | '.join(row)}",
                                    page=table.page_number,
                                    section=table.section_name,
                                    block_type="table",
                                    confidence=0.90
                                )
                                extractions.append(("annual_turnover", value, evidence))
                                break
                            except ValueError:
                                continue
        
        return extractions
    
    def _evaluate_criteria(
        self,
        extraction_results: List[ExtractionResult]
    ) -> List[CriterionEvaluation]:
        """Evaluate tender criteria against extracted data"""
        evaluations = []
        
        for criterion in self.tender_criteria:
            field = criterion.get('field', '')
            
            # Find extraction result for this field
            extraction = next((r for r in extraction_results if r.field == field), None)
            
            if extraction:
                evaluation = evaluate_criterion(criterion, extraction)
                evaluations.append(evaluation)
        
        return evaluations
    
    def _format_extraction_result(self, result: ExtractionResult) -> Dict:
        """Format extraction result for JSON output"""
        return {
            "field": result.field,
            "value": result.value,
            "raw_text": result.raw_text,
            "confidence": result.confidence,
            "extraction_method": result.extraction_method,
            "status": result.status.value,
            "evidence": asdict(result.evidence) if result.evidence else None,
            "note": result.note,
        }
    
    def _format_criterion_evaluation(self, evaluation: CriterionEvaluation) -> Dict:
        """Format criterion evaluation for JSON output"""
        return {
            "criterion": evaluation.criterion_label,
            "criterion_id": evaluation.criterion_id,
            "required_value": evaluation.required_value,
            "extracted_value": evaluation.extracted_value,
            "meets_requirement": evaluation.meets_requirement,
            "confidence": evaluation.confidence,
            "status": evaluation.status.value,
            "evidence": asdict(evaluation.evidence) if evaluation.evidence else None,
            "reasoning": evaluation.reasoning,
        }
    
    def _convert_to_legacy_format(
        self,
        extraction_results: List[ExtractionResult]
    ) -> List[Dict]:
        """Convert new format to legacy ExtractedField format for backward compatibility"""
        legacy_fields = []
        
        for result in extraction_results:
            if result.value is not None:
                legacy_fields.append({
                    "field": result.field,
                    "value": result.value,
                    "raw_text": result.raw_text,
                    "confidence": result.confidence,
                    "source_page": result.evidence.page if result.evidence else 1,
                    "source_file": None,
                    "needs_review": result.status in (ExtractionStatus.CONFLICT, ExtractionStatus.LOW_CONFIDENCE),
                    "note": result.note,
                })
        
        return legacy_fields


# ── Convenience function for API ─────────────────────────────────────────────

def process_bidder_files(
    file_paths: List[str],
    bidder_name: str,
    tender_id: str,
    tender_criteria: Optional[List[Dict]] = None
) -> BidderData:
    """
    Process bidder files with production-grade pipeline
    
    Args:
        file_paths: List of file paths
        bidder_name: Bidder company name
        tender_id: Tender ID
        tender_criteria: Optional list of criteria to evaluate
    
    Returns:
        BidderData object
    """
    processor = BidderProcessor(tender_criteria=tender_criteria)
    data = processor.process_bidder_files(
        file_paths=file_paths,
        bidder_name=bidder_name,
        tender_id=tender_id
    )
    
    # Convert to BidderData model
    return BidderData(**data)
